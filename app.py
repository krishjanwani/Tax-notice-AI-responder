import streamlit as st
import os
from google import genai
from dotenv import load_dotenv
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

# 1. PAGE CONFIGURATION
st.set_page_config(page_title="CA AI Assistant", layout="centered")
st.title("⚖️ Tax Notice AI-Responder")
st.caption("Developed for Professional CA Practice - Krish Janwani & Associates")

# 2. SECURE API KEY HANDLING
# Prioritizes Streamlit Cloud Secrets for the live link
if "GEMINI_API_KEY" in st.secrets:
    final_key = st.secrets["GEMINI_API_KEY"]
else:
    load_dotenv()
    final_key = os.getenv("GEMINI_API_KEY")

# 3. CLEAN WORD DOCUMENT FUNCTION
def create_docx(text):
    doc = Document()
    
    # Add a professional centered heading
    header = doc.add_heading('DRAFT RESPONSE TO INCOME TAX NOTICE', 0)
    header.alignment = 1 
    
    doc.add_paragraph("From: Krish Janwani & Associates (Chartered Accountants)")
    doc.add_paragraph("-" * 60)
    
    # --- MARKDOWN TO WORD CONVERSION LOGIC ---
    # This removes messy '**' and converts them to actual BOLD text
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph()
            # Split by double asterisks to find bold segments
            parts = para_text.split('**')
            for i, part in enumerate(parts):
                # Clean up other markdown symbols like ### or ---
                clean_part = part.replace('###', '').replace('---', '')
                if i % 2 == 1:
                    # Odd-indexed parts were inside **...**, so make them bold
                    p.add_run(clean_part).bold = True
                else:
                    p.add_run(clean_part)
    
    # Save to memory buffer
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 4. USER INTERFACE
file = st.file_uploader("Upload Tax Notice (PDF)", type="pdf")

if file and final_key:
    try:
        # Extract text from PDF
        pdf_data = fitz.open(stream=file.read(), filetype="pdf")
        notice_text = "".join([page.get_text() for page in pdf_data])
        st.info("Notice content loaded. Ready to generate.")
    except Exception as e:
        st.error(f"Error reading PDF: {e}")

    if st.button("Generate Professional Draft"):
        try:
            client = genai.Client(api_key=final_key)
            
            with st.spinner("AI is analyzing notice and drafting response..."):
                response = client.models.generate_content(
                    model="gemini-2.5-flash", 
                    contents=f"Act as a senior CA. Draft a formal, clean reply to this tax notice: {notice_text}"
                )
            
            draft_content = response.text
            
            # Show preview
            st.subheader("Draft Preview")
            st.markdown(draft_content)
            
            # --- CLEAN WORD DOWNLOAD ---
            st.divider()
            docx_data = create_docx(draft_content)
            st.download_button(
                label="📥 Download Clean Word Document (.docx)",
                data=docx_data,
                file_name="Clean_Tax_Response.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
                
        except Exception as e:
            if "429" in str(e):
                st.error("Quota Exceeded: Please wait 60 seconds and try again.")
            elif "403" in str(e):
                st.error("API Key Error: Please check your Streamlit Secrets for a valid key.")
            else:
                st.error(f"System Error: {e}")
            
elif not final_key:
    st.warning("Please provide a Gemini API Key to enable AI features.")
